from __future__ import annotations

import asyncio
import csv
import json
import re
from dataclasses import dataclass
from datetime import UTC, datetime, time, timedelta
from datetime import date as dt_date
from difflib import SequenceMatcher
from io import StringIO
from pathlib import Path
from typing import Any, Literal
from uuid import UUID, uuid4
from zoneinfo import ZoneInfo

from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile, status
from openpyxl import load_workbook
from pydantic import BaseModel, ConfigDict, Field, ValidationError
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.config import get_settings
from app.core.logging import get_logger, sanitize_log_data
from app.core.time import utc_now
from app.db.enums import AIImportDraftStatus, AIImportMode, LessonStatus, RoleCode
from app.db.models import AIImportDraft, Discipline, Faculty, Group, Stream, User
from app.integrations.openrouter_client import OpenRouterError, openrouter_chat_completion
from app.services.audit import log_audit
from app.services.import_apply import (
    ensure_student_membership,
    ensure_teacher_assignment,
    resolve_discipline,
    resolve_faculty,
    resolve_group,
    resolve_stream,
    resolve_user,
    upsert_lesson,
)

settings = get_settings()
logger = get_logger(__name__)

MappingAction = Literal["match_existing", "create_new", "unresolved"]
IssueSeverity = Literal["error", "warning", "info"]
WeekParity = Literal["all", "odd", "even"]
DocumentKind = Literal["mixed", "users", "schedule"]

SUPPORTED_AI_IMPORT_EXTENSIONS = {".csv", ".xlsx", ".pdf", ".docx"}
SUPPORTED_AI_IMPORT_MIME_TYPES = {
    "text/csv",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _draft_row_id() -> str:
    return uuid4().hex


class AIImportWizard(BaseModel):
    term_start: dt_date | None = None
    term_end: dt_date | None = None
    first_week_parity: Literal["odd", "even"] | None = None


class DraftRowBase(BaseModel):
    model_config = ConfigDict(extra="ignore")

    draft_id: str = Field(default_factory=_draft_row_id)
    action: MappingAction = "unresolved"
    existing_id: UUID | None = None
    confidence: float | None = None
    source_ref: str | None = None


class DraftFacultyRow(DraftRowBase):
    code: str | None = None
    name: str | None = None


class DraftStreamRow(DraftRowBase):
    name: str | None = None
    faculty_code: str | None = None


class DraftGroupRow(DraftRowBase):
    code: str | None = None
    name: str | None = None
    faculty_code: str | None = None
    stream_name: str | None = None
    parent_group_code: str | None = None
    is_subgroup: bool = False


class DraftDisciplineRow(DraftRowBase):
    code: str | None = None
    name: str | None = None


class DraftUserRow(DraftRowBase):
    username: str | None = None
    full_name: str | None = None
    email: str | None = None
    roles: list[str] = Field(default_factory=list)
    group_code: str | None = None


class DraftMembershipRow(BaseModel):
    model_config = ConfigDict(extra="ignore")

    draft_id: str = Field(default_factory=_draft_row_id)
    student_username: str | None = None
    student_full_name: str | None = None
    group_code: str | None = None
    start_date: dt_date | None = None
    source_ref: str | None = None


class DraftAssignmentRow(BaseModel):
    model_config = ConfigDict(extra="ignore")

    draft_id: str = Field(default_factory=_draft_row_id)
    teacher_username: str | None = None
    teacher_full_name: str | None = None
    discipline_code: str | None = None
    discipline_name: str | None = None
    group_code: str | None = None
    source_ref: str | None = None


class DraftSchedulePatternRow(BaseModel):
    model_config = ConfigDict(extra="ignore")

    draft_id: str = Field(default_factory=_draft_row_id)
    group_code: str | None = None
    discipline_code: str | None = None
    discipline_name: str | None = None
    teacher_username: str | None = None
    teacher_name: str | None = None
    date: dt_date | None = None
    day_of_week: str | None = None
    start_time: str | None = None
    end_time: str | None = None
    week_parity: WeekParity = "all"
    room: str | None = None
    note: str | None = None
    source_ref: str | None = None


class DraftLessonRow(BaseModel):
    model_config = ConfigDict(extra="ignore")

    draft_id: str = Field(default_factory=_draft_row_id)
    pattern_draft_id: str | None = None
    group_code: str | None = None
    discipline_code: str | None = None
    discipline_name: str | None = None
    teacher_username: str | None = None
    teacher_name: str | None = None
    starts_at: datetime
    ends_at: datetime
    room: str | None = None
    status: str = LessonStatus.PLANNED.value
    source_ref: str | None = None


class DraftEntities(BaseModel):
    model_config = ConfigDict(extra="ignore")

    faculties: list[DraftFacultyRow] = Field(default_factory=list)
    streams: list[DraftStreamRow] = Field(default_factory=list)
    groups: list[DraftGroupRow] = Field(default_factory=list)
    disciplines: list[DraftDisciplineRow] = Field(default_factory=list)
    users: list[DraftUserRow] = Field(default_factory=list)
    memberships: list[DraftMembershipRow] = Field(default_factory=list)
    assignments: list[DraftAssignmentRow] = Field(default_factory=list)


class DraftPayload(BaseModel):
    model_config = ConfigDict(extra="ignore")

    detected_doc_kind: DocumentKind = "mixed"
    notes: list[str] = Field(default_factory=list)
    entities: DraftEntities = Field(default_factory=DraftEntities)
    schedule_patterns: list[DraftSchedulePatternRow] = Field(default_factory=list)
    lessons: list[DraftLessonRow] = Field(default_factory=list)


class DraftIssue(BaseModel):
    model_config = ConfigDict(extra="ignore")

    severity: IssueSeverity
    code: str
    message: str
    source_ref: str | None = None
    field_path: str | None = None
    requires_action: bool = False


class ParsedFaculty(BaseModel):
    model_config = ConfigDict(extra="ignore")

    code: str | None = None
    name: str | None = None
    source_ref: str | None = None


class ParsedStream(BaseModel):
    model_config = ConfigDict(extra="ignore")

    name: str | None = None
    faculty_code: str | None = None
    source_ref: str | None = None


class ParsedGroup(BaseModel):
    model_config = ConfigDict(extra="ignore")

    code: str | None = None
    name: str | None = None
    faculty_code: str | None = None
    stream_name: str | None = None
    parent_group_code: str | None = None
    is_subgroup: bool = False
    source_ref: str | None = None


class ParsedDiscipline(BaseModel):
    model_config = ConfigDict(extra="ignore")

    code: str | None = None
    name: str | None = None
    source_ref: str | None = None


class ParsedUser(BaseModel):
    model_config = ConfigDict(extra="ignore")

    username: str | None = None
    full_name: str | None = None
    email: str | None = None
    roles: list[str] = Field(default_factory=list)
    group_code: str | None = None
    source_ref: str | None = None


class ParsedSchedulePattern(BaseModel):
    model_config = ConfigDict(extra="ignore")

    group_code: str | None = None
    discipline_code: str | None = None
    discipline_name: str | None = None
    teacher_username: str | None = None
    teacher_name: str | None = None
    date: dt_date | None = None
    day_of_week: str | None = None
    start_time: str | None = None
    end_time: str | None = None
    week_parity: WeekParity = "all"
    room: str | None = None
    note: str | None = None
    source_ref: str | None = None


class ParsedDocument(BaseModel):
    model_config = ConfigDict(extra="ignore")

    detected_doc_kind: DocumentKind = "mixed"
    notes: list[str] = Field(default_factory=list)
    faculties: list[ParsedFaculty] = Field(default_factory=list)
    streams: list[ParsedStream] = Field(default_factory=list)
    groups: list[ParsedGroup] = Field(default_factory=list)
    disciplines: list[ParsedDiscipline] = Field(default_factory=list)
    users: list[ParsedUser] = Field(default_factory=list)
    schedule_patterns: list[ParsedSchedulePattern] = Field(default_factory=list)


@dataclass(slots=True)
class ExtractedBlock:
    source_ref: str
    text: str


@dataclass(slots=True)
class ExtractedDocument:
    file_name: str
    file_ext: str
    blocks: list[ExtractedBlock]
    metadata: dict[str, Any]


def ai_import_root() -> Path:
    preferred = Path(settings.ai_imports_dir)
    try:
        preferred.mkdir(parents=True, exist_ok=True)
        return preferred
    except OSError:
        fallback = Path("/tmp/universe-ai-imports")
        fallback.mkdir(parents=True, exist_ok=True)
        return fallback


def ai_import_sources_dir() -> Path:
    path = ai_import_root() / "sources"
    path.mkdir(parents=True, exist_ok=True)
    return path


def _ai_import_excerpt(blocks: list[ExtractedBlock]) -> str:
    combined = "\n\n".join(block.text for block in blocks[:4]).strip()
    if len(combined) > 600:
        return f"{combined[:600]}...[truncated]"
    return combined


def _coerce_non_empty(value: str | None) -> str | None:
    if value is None:
        return None
    normalized = re.sub(r"\s+", " ", str(value)).strip()
    return normalized or None


def _normalize_code(value: str | None) -> str | None:
    value = _coerce_non_empty(value)
    return value.upper() if value else None


def _normalize_day_of_week(value: str | None) -> str | None:
    value = _coerce_non_empty(value)
    if not value:
        return None
    mapping = {
        "monday": "monday",
        "понедельник": "monday",
        "mon": "monday",
        "tuesday": "tuesday",
        "вторник": "tuesday",
        "tue": "tuesday",
        "wednesday": "wednesday",
        "среда": "wednesday",
        "wed": "wednesday",
        "thursday": "thursday",
        "четверг": "thursday",
        "thu": "thursday",
        "friday": "friday",
        "пятница": "friday",
        "fri": "friday",
        "saturday": "saturday",
        "суббота": "saturday",
        "sat": "saturday",
        "sunday": "sunday",
        "воскресенье": "sunday",
        "sun": "sunday",
    }
    return mapping.get(value.lower(), value.lower())


def _normalize_time(value: str | None) -> str | None:
    value = _coerce_non_empty(value)
    if not value:
        return None
    value = value.replace(".", ":")
    match = re.fullmatch(r"(\d{1,2}):(\d{2})", value)
    if not match:
        return None
    return f"{int(match.group(1)):02d}:{match.group(2)}"


def _parse_time_value(value: str) -> time:
    parsed = _normalize_time(value)
    if not parsed:
        raise ValueError(f"Invalid time value '{value}'")
    hours, minutes = parsed.split(":")
    return time(hour=int(hours), minute=int(minutes))


def _normalize_week_parity(value: str | None) -> WeekParity:
    value = _coerce_non_empty(value)
    if not value:
        return "all"
    normalized = value.lower()
    if normalized in {"all", "every", "каждую", "каждая", "все", "weekly"}:
        return "all"
    if normalized in {"odd", "нечет", "нечетная", "нечёт", "нечётная"}:
        return "odd"
    if normalized in {"even", "чет", "четная", "чёт", "чётная"}:
        return "even"
    return "all"


def _normalize_roles(raw_roles: list[str]) -> list[str]:
    mapping = {
        "student": "student",
        "студент": "student",
        "teacher": "teacher",
        "преподаватель": "teacher",
        "admin": "admin",
        "administrator": "admin",
        "админ": "admin",
        "curator": "curator",
        "куратор": "curator",
        "тьютор": "curator",
    }
    normalized: list[str] = []
    for role in raw_roles:
        value = _coerce_non_empty(role)
        if not value:
            continue
        mapped = mapping.get(value.lower())
        if mapped and mapped not in normalized:
            normalized.append(mapped)
    return normalized


def _role_codes_from_strings(raw_roles: list[str]) -> list[RoleCode]:
    mapping = {
        "student": RoleCode.STUDENT,
        "teacher": RoleCode.TEACHER,
        "admin": RoleCode.ADMIN,
        "curator": RoleCode.CURATOR,
    }
    return [mapping[role] for role in raw_roles if role in mapping]


def _sequence_score(left: str | None, right: str | None) -> float:
    if not left or not right:
        return 0.0
    return SequenceMatcher(a=left.lower(), b=right.lower()).ratio()


def _serialize_table(rows: list[list[str | None]]) -> str:
    if not rows:
        return ""
    buffer = StringIO()
    writer = csv.writer(buffer)
    for row in rows:
        writer.writerow([_coerce_non_empty(item) or "" for item in row])
    return buffer.getvalue().strip()


def _trim_blocks(blocks: list[ExtractedBlock]) -> list[ExtractedBlock]:
    trimmed: list[ExtractedBlock] = []
    total_chars = 0
    for block in blocks:
        text = block.text.strip()
        if not text:
            continue
        remaining = settings.ai_import_max_chars - total_chars
        if remaining <= 0:
            break
        if len(text) > remaining:
            text = f"{text[:remaining]}...[truncated]"
        trimmed.append(ExtractedBlock(source_ref=block.source_ref, text=text))
        total_chars += len(text)
    return trimmed


async def save_ai_import_source(file: UploadFile) -> tuple[str, str]:
    ext = Path(file.filename or "").suffix.lower()
    if ext not in SUPPORTED_AI_IMPORT_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only csv/xlsx/pdf/docx files are supported",
        )
    if file.content_type and file.content_type not in SUPPORTED_AI_IMPORT_MIME_TYPES:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unsupported AI import file type")

    content = await file.read()
    max_bytes = settings.ai_import_max_size_mb * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="AI import file exceeds size limit")

    sources_dir = await asyncio.to_thread(ai_import_sources_dir)
    path = sources_dir / f"{uuid4().hex}{ext}"
    await asyncio.to_thread(path.write_bytes, content)
    return file.filename or path.name, str(path)


def _extract_csv_document(file_path: str) -> ExtractedDocument:
    path = Path(file_path)
    with path.open("r", encoding="utf-8", newline="") as handle:
        reader = csv.reader(handle)
        rows = [row for row in reader][:250]
    blocks = [ExtractedBlock(source_ref="csv:sheet1", text=_serialize_table(rows))]
    return ExtractedDocument(
        file_name=path.name,
        file_ext=path.suffix.lower(),
        blocks=_trim_blocks(blocks),
        metadata={"rows": len(rows), "sheets": 1},
    )


def _extract_xlsx_document(file_path: str) -> ExtractedDocument:
    path = Path(file_path)
    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[ExtractedBlock] = []
    sheet_count = 0
    try:
        for worksheet in workbook.worksheets[: settings.ai_import_max_pages]:
            sheet_count += 1
            rows: list[list[str | None]] = []
            for index, row in enumerate(worksheet.iter_rows(values_only=True), start=1):
                rows.append([str(cell) if cell is not None else None for cell in row])
                if index >= 250:
                    break
            blocks.append(
                ExtractedBlock(
                    source_ref=f"xlsx:{worksheet.title}",
                    text=_serialize_table(rows),
                )
            )
    finally:
        workbook.close()
    return ExtractedDocument(
        file_name=path.name,
        file_ext=path.suffix.lower(),
        blocks=_trim_blocks(blocks),
        metadata={"sheets": sheet_count},
    )


def _extract_docx_document(file_path: str) -> ExtractedDocument:
    path = Path(file_path)
    document = DocxDocument(path)
    blocks: list[ExtractedBlock] = []
    paragraph_text = "\n".join(_coerce_non_empty(paragraph.text) or "" for paragraph in document.paragraphs)
    if paragraph_text.strip():
        blocks.append(ExtractedBlock(source_ref="docx:paragraphs", text=paragraph_text.strip()))
    for index, table in enumerate(document.tables[: settings.ai_import_max_pages], start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        blocks.append(ExtractedBlock(source_ref=f"docx:table:{index}", text=_serialize_table(rows)))
    return ExtractedDocument(
        file_name=path.name,
        file_ext=path.suffix.lower(),
        blocks=_trim_blocks(blocks),
        metadata={"paragraphs": len(document.paragraphs), "tables": len(document.tables)},
    )


def _extract_pdf_document(file_path: str) -> ExtractedDocument:
    path = Path(file_path)
    import pdfplumber
    from pypdf import PdfReader

    blocks: list[ExtractedBlock] = []
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages[: settings.ai_import_max_pages], start=1):
            text = _coerce_non_empty(page.extract_text())
            if text:
                blocks.append(ExtractedBlock(source_ref=f"pdf:page:{index}", text=text))
            for table_index, table in enumerate(page.extract_tables() or [], start=1):
                blocks.append(
                    ExtractedBlock(
                        source_ref=f"pdf:page:{index}:table:{table_index}",
                        text=_serialize_table(table),
                    )
                )

    if not blocks:
        reader = PdfReader(path)
        for index, page in enumerate(reader.pages[: settings.ai_import_max_pages], start=1):
            text = _coerce_non_empty(page.extract_text())
            if text:
                blocks.append(ExtractedBlock(source_ref=f"pdf:page:{index}", text=text))

    return ExtractedDocument(
        file_name=path.name,
        file_ext=path.suffix.lower(),
        blocks=_trim_blocks(blocks),
        metadata={"pages": len(blocks)},
    )


def extract_document(file_path: str) -> ExtractedDocument:
    ext = Path(file_path).suffix.lower()
    if ext == ".csv":
        return _extract_csv_document(file_path)
    if ext == ".xlsx":
        return _extract_xlsx_document(file_path)
    if ext == ".docx":
        return _extract_docx_document(file_path)
    if ext == ".pdf":
        return _extract_pdf_document(file_path)
    raise ValueError(f"Unsupported AI import extension '{ext}'")


def _llm_prompt(mode: AIImportMode, extracted: ExtractedDocument) -> str:
    blocks = "\n\n".join(
        f"[{block.source_ref}]\n{block.text}" for block in extracted.blocks if block.text.strip()
    )
    return (
        "Ты помогаешь импортировать академические данные в систему учета посещаемости.\n"
        "Нужно извлечь только факты из документа и вернуть строго JSON без Markdown и без пояснений.\n"
        "Не выдумывай коды, логины, даты или роли. Если данных нет, используй null или пустой список.\n"
        "Для расписания извлекай weekly patterns: group_code, discipline_code/discipline_name, "
        "teacher_username/teacher_name, date или day_of_week, start_time, end_time, week_parity, room, note.\n"
        "Для пользователей roles используй только student, teacher, admin, curator.\n"
        "Если документ смешанный, detected_doc_kind должен быть mixed.\n"
        "Верни JSON по схеме:\n"
        "{"
        '"detected_doc_kind":"mixed|users|schedule",'
        '"notes":["..."],'
        '"faculties":[{"code":null,"name":null,"source_ref":"..."}],'
        '"streams":[{"name":null,"faculty_code":null,"source_ref":"..."}],'
        '"groups":[{"code":null,"name":null,"faculty_code":null,"stream_name":null,"parent_group_code":null,"is_subgroup":false,"source_ref":"..."}],'
        '"disciplines":[{"code":null,"name":null,"source_ref":"..."}],'
        '"users":[{"username":null,"full_name":null,"email":null,"roles":["student"],"group_code":null,"source_ref":"..."}],'
        '"schedule_patterns":[{"group_code":null,"discipline_code":null,"discipline_name":null,"teacher_username":null,"teacher_name":null,"date":null,"day_of_week":null,"start_time":null,"end_time":null,"week_parity":"all","room":null,"note":null,"source_ref":"..."}]'
        "}\n"
        f"Выбранный режим импорта: {mode.value}.\n"
        f"Имя файла: {extracted.file_name}.\n"
        "Документ:\n"
        f"{blocks or '[empty]'}"
    )


def _extract_json_from_text(text: str) -> str:
    stripped = text.strip()
    if stripped.startswith("```"):
        match = re.search(r"```(?:json)?\s*(\{.*\})\s*```", stripped, flags=re.S)
        if match:
            return match.group(1)
    start = stripped.find("{")
    end = stripped.rfind("}")
    if start == -1 or end == -1 or end < start:
        raise ValueError("LLM response did not contain JSON")
    return stripped[start : end + 1]


async def _normalize_with_llm(mode: AIImportMode, extracted: ExtractedDocument) -> ParsedDocument:
    if not settings.openrouter_api_key:
        raise RuntimeError("OPENROUTER_API_KEY is not configured")

    messages = [
        {
            "role": "system",
            "content": (
                "Ты извлекаешь данные из учебных документов и возвращаешь только корректный JSON. "
                "Никакого Markdown, никаких комментариев."
            ),
        },
        {"role": "user", "content": _llm_prompt(mode, extracted)},
    ]
    try:
        text = await openrouter_chat_completion(
            messages,
            model=settings.ai_import_model,
            temperature=0,
            timeout_seconds=settings.ai_import_timeout_seconds,
        )
    except OpenRouterError as exc:
        logger.warning(
            "ai_import_llm_failed",
            file_name=extracted.file_name,
            metadata=sanitize_log_data(extracted.metadata),
            reason=str(exc),
        )
        raise RuntimeError("AI normalization request failed") from exc

    try:
        payload = json.loads(_extract_json_from_text(text))
        return ParsedDocument.model_validate(payload)
    except (json.JSONDecodeError, ValidationError, ValueError) as exc:
        raise RuntimeError("AI normalization returned invalid JSON") from exc


def _dedupe_rows[T: BaseModel](rows: list[T], key_builder) -> list[T]:
    unique: dict[str, T] = {}
    for row in rows:
        key = key_builder(row)
        if not key:
            unique[row.model_dump_json()] = row
            continue
        if key not in unique:
            unique[key] = row
    return list(unique.values())


def _parsed_to_payload(parsed: ParsedDocument) -> DraftPayload:
    payload = DraftPayload(
        detected_doc_kind=parsed.detected_doc_kind,
        notes=[_coerce_non_empty(note) or "" for note in parsed.notes if _coerce_non_empty(note)],
        entities=DraftEntities(
            faculties=[
                DraftFacultyRow(
                    code=_normalize_code(row.code),
                    name=_coerce_non_empty(row.name) or _normalize_code(row.code),
                    source_ref=row.source_ref,
                )
                for row in parsed.faculties
            ],
            streams=[
                DraftStreamRow(
                    name=_coerce_non_empty(row.name),
                    faculty_code=_normalize_code(row.faculty_code),
                    source_ref=row.source_ref,
                )
                for row in parsed.streams
            ],
            groups=[
                DraftGroupRow(
                    code=_normalize_code(row.code),
                    name=_coerce_non_empty(row.name) or _normalize_code(row.code),
                    faculty_code=_normalize_code(row.faculty_code),
                    stream_name=_coerce_non_empty(row.stream_name),
                    parent_group_code=_normalize_code(row.parent_group_code),
                    is_subgroup=row.is_subgroup,
                    source_ref=row.source_ref,
                )
                for row in parsed.groups
            ],
            disciplines=[
                DraftDisciplineRow(
                    code=_normalize_code(row.code),
                    name=_coerce_non_empty(row.name) or _normalize_code(row.code),
                    source_ref=row.source_ref,
                )
                for row in parsed.disciplines
            ],
            users=[
                DraftUserRow(
                    username=_coerce_non_empty(row.username),
                    full_name=_coerce_non_empty(row.full_name) or _coerce_non_empty(row.username),
                    email=_coerce_non_empty(row.email),
                    roles=_normalize_roles(row.roles),
                    group_code=_normalize_code(row.group_code),
                    source_ref=row.source_ref,
                )
                for row in parsed.users
            ],
        ),
        schedule_patterns=[
            DraftSchedulePatternRow(
                group_code=_normalize_code(row.group_code),
                discipline_code=_normalize_code(row.discipline_code),
                discipline_name=_coerce_non_empty(row.discipline_name),
                teacher_username=_coerce_non_empty(row.teacher_username),
                teacher_name=_coerce_non_empty(row.teacher_name),
                date=row.date,
                day_of_week=_normalize_day_of_week(row.day_of_week),
                start_time=_normalize_time(row.start_time),
                end_time=_normalize_time(row.end_time),
                week_parity=_normalize_week_parity(row.week_parity),
                room=_coerce_non_empty(row.room),
                note=_coerce_non_empty(row.note),
                source_ref=row.source_ref,
            )
            for row in parsed.schedule_patterns
        ],
    )
    return _dedupe_payload(payload)


def _dedupe_payload(payload: DraftPayload) -> DraftPayload:
    payload.entities.faculties = _dedupe_rows(payload.entities.faculties, lambda row: row.code or row.name or "")
    payload.entities.streams = _dedupe_rows(
        payload.entities.streams,
        lambda row: f"{row.faculty_code or ''}:{row.name or ''}",
    )
    payload.entities.groups = _dedupe_rows(payload.entities.groups, lambda row: row.code or row.name or "")
    payload.entities.disciplines = _dedupe_rows(
        payload.entities.disciplines,
        lambda row: row.code or row.name or "",
    )
    payload.entities.users = _dedupe_rows(
        payload.entities.users,
        lambda row: row.username or row.email or row.full_name or "",
    )
    payload.schedule_patterns = _dedupe_rows(
        payload.schedule_patterns,
        lambda row: "|".join(
            [
                row.group_code or "",
                row.discipline_code or row.discipline_name or "",
                row.teacher_username or row.teacher_name or "",
                row.date.isoformat() if row.date else row.day_of_week or "",
                row.start_time or "",
                row.end_time or "",
                row.week_parity,
            ]
        ),
    )
    return payload


def _derive_memberships_and_assignments(payload: DraftPayload) -> None:
    memberships = payload.entities.memberships
    assignments = payload.entities.assignments

    for user in payload.entities.users:
        if "student" in user.roles and user.group_code:
            memberships.append(
                DraftMembershipRow(
                    student_username=user.username,
                    student_full_name=user.full_name,
                    group_code=user.group_code,
                    source_ref=user.source_ref,
                )
            )

    for pattern in payload.schedule_patterns:
        if pattern.group_code and not any(group.code == pattern.group_code for group in payload.entities.groups):
            payload.entities.groups.append(
                DraftGroupRow(
                    code=pattern.group_code,
                    name=pattern.group_code,
                    source_ref=pattern.source_ref,
                    action="create_new",
                    confidence=0.35,
                )
            )
        discipline_key = pattern.discipline_code or pattern.discipline_name
        if discipline_key and not any(
            (row.code and row.code == pattern.discipline_code)
            or (row.name and row.name == pattern.discipline_name)
            for row in payload.entities.disciplines
        ):
            payload.entities.disciplines.append(
                DraftDisciplineRow(
                    code=pattern.discipline_code,
                    name=pattern.discipline_name or pattern.discipline_code,
                    source_ref=pattern.source_ref,
                    action="create_new" if pattern.discipline_code else "unresolved",
                    confidence=0.35 if pattern.discipline_code else None,
                )
            )
        if (pattern.teacher_username or pattern.teacher_name) and not any(
            (row.username and row.username == pattern.teacher_username)
            or (row.full_name and row.full_name == pattern.teacher_name)
            for row in payload.entities.users
        ):
            payload.entities.users.append(
                DraftUserRow(
                    username=pattern.teacher_username,
                    full_name=pattern.teacher_name or pattern.teacher_username,
                    roles=["teacher"],
                    source_ref=pattern.source_ref,
                    action="unresolved",
                )
            )
        assignments.append(
            DraftAssignmentRow(
                teacher_username=pattern.teacher_username,
                teacher_full_name=pattern.teacher_name,
                discipline_code=pattern.discipline_code,
                discipline_name=pattern.discipline_name,
                group_code=pattern.group_code,
                source_ref=pattern.source_ref,
            )
        )

    payload.entities.memberships = _dedupe_rows(
        memberships,
        lambda row: f"{row.student_username or row.student_full_name or ''}:{row.group_code or ''}",
    )
    payload.entities.assignments = _dedupe_rows(
        assignments,
        lambda row: "|".join(
            [
                row.teacher_username or row.teacher_full_name or "",
                row.discipline_code or row.discipline_name or "",
                row.group_code or "",
            ]
        ),
    )
    _dedupe_payload(payload)


async def _load_catalog(session: AsyncSession) -> dict[str, Any]:
    faculties = (await session.execute(select(Faculty))).scalars().all()
    streams = (await session.execute(select(Stream))).scalars().all()
    groups = (await session.execute(select(Group))).scalars().all()
    disciplines = (await session.execute(select(Discipline))).scalars().all()
    users = (await session.execute(select(User).options(selectinload(User.roles)))).scalars().all()
    return {
        "faculties": faculties,
        "streams": streams,
        "groups": groups,
        "disciplines": disciplines,
        "users": users,
    }


def _set_manual_mapping_defaults[T: DraftRowBase](row: T) -> T:
    if row.action in {"match_existing", "create_new"}:
        return row
    row.action = "unresolved"
    row.existing_id = None
    row.confidence = None
    return row


def _find_best_match_by_code_or_name(
    rows: list[Any],
    *,
    code: str | None,
    name: str | None,
    code_getter,
    name_getter,
) -> tuple[Any | None, float, bool]:
    if code:
        exact = [row for row in rows if _normalize_code(code_getter(row)) == _normalize_code(code)]
        if len(exact) == 1:
            return exact[0], 0.99, False
        if len(exact) > 1:
            return None, 0.0, True
    if name:
        exact = [row for row in rows if _coerce_non_empty(name_getter(row)) == _coerce_non_empty(name)]
        if len(exact) == 1:
            return exact[0], 0.92, False
        if len(exact) > 1:
            return None, 0.0, True
        fuzzy = sorted(
            ((row, _sequence_score(name_getter(row), name)) for row in rows),
            key=lambda item: item[1],
            reverse=True,
        )
        if fuzzy and fuzzy[0][1] >= 0.88:
            if len(fuzzy) > 1 and abs(fuzzy[0][1] - fuzzy[1][1]) < 0.03:
                return None, 0.0, True
            return fuzzy[0][0], fuzzy[0][1], False
    return None, 0.0, False


def _apply_auto_mapping(row: DraftRowBase, *, matched_id: UUID | None, confidence: float | None) -> None:
    if matched_id:
        row.action = "match_existing"
        row.existing_id = matched_id
        row.confidence = confidence
    elif row.action == "unresolved":
        row.existing_id = None
        row.confidence = confidence


async def _match_payload(session: AsyncSession, payload: DraftPayload) -> DraftPayload:
    catalog = await _load_catalog(session)

    for row in payload.entities.faculties:
        _set_manual_mapping_defaults(row)
        if row.action == "match_existing" and row.existing_id:
            row.confidence = row.confidence or 1.0
            continue
        if row.action == "create_new":
            row.confidence = row.confidence or 0.4
            continue
        matched, confidence, ambiguous = _find_best_match_by_code_or_name(
            catalog["faculties"],
            code=row.code,
            name=row.name,
            code_getter=lambda item: item.code,
            name_getter=lambda item: item.name,
        )
        if matched:
            _apply_auto_mapping(row, matched_id=matched.id, confidence=confidence)
        elif not ambiguous and row.code and row.name and row.action == "unresolved":
            row.action = "create_new"
            row.confidence = 0.4

    for row in payload.entities.streams:
        _set_manual_mapping_defaults(row)
        if row.action == "match_existing" and row.existing_id:
            row.confidence = row.confidence or 1.0
            continue
        if row.action == "create_new":
            row.confidence = row.confidence or 0.4
            continue
        candidates = catalog["streams"]
        if row.faculty_code:
            faculty = next(
                (item for item in payload.entities.faculties if item.code == row.faculty_code and item.existing_id),
                None,
            )
            if faculty:
                candidates = [stream for stream in candidates if stream.faculty_id == faculty.existing_id]
        matched, confidence, ambiguous = _find_best_match_by_code_or_name(
            candidates,
            code=None,
            name=row.name,
            code_getter=lambda _item: None,
            name_getter=lambda item: item.name,
        )
        if matched:
            _apply_auto_mapping(row, matched_id=matched.id, confidence=confidence)
        elif not ambiguous and row.name and row.faculty_code and row.action == "unresolved":
            row.action = "create_new"
            row.confidence = 0.4

    for row in payload.entities.groups:
        _set_manual_mapping_defaults(row)
        if row.action == "match_existing" and row.existing_id:
            row.confidence = row.confidence or 1.0
            continue
        if row.action == "create_new":
            row.confidence = row.confidence or 0.4
            continue
        matched, confidence, ambiguous = _find_best_match_by_code_or_name(
            catalog["groups"],
            code=row.code,
            name=row.name,
            code_getter=lambda item: item.code,
            name_getter=lambda item: item.name,
        )
        if matched:
            _apply_auto_mapping(row, matched_id=matched.id, confidence=confidence)
        elif not ambiguous and row.code and row.name and row.action == "unresolved":
            row.action = "create_new"
            row.confidence = 0.4

    for row in payload.entities.disciplines:
        _set_manual_mapping_defaults(row)
        if row.action == "match_existing" and row.existing_id:
            row.confidence = row.confidence or 1.0
            continue
        if row.action == "create_new":
            row.confidence = row.confidence or 0.4
            continue
        matched, confidence, ambiguous = _find_best_match_by_code_or_name(
            catalog["disciplines"],
            code=row.code,
            name=row.name,
            code_getter=lambda item: item.code,
            name_getter=lambda item: item.name,
        )
        if matched:
            _apply_auto_mapping(row, matched_id=matched.id, confidence=confidence)
        elif not ambiguous and row.code and row.name and row.action == "unresolved":
            row.action = "create_new"
            row.confidence = 0.4

    for row in payload.entities.users:
        _set_manual_mapping_defaults(row)
        if row.action == "match_existing" and row.existing_id:
            row.confidence = row.confidence or 1.0
            continue
        if row.action == "create_new":
            row.confidence = row.confidence or 0.4
            continue
        candidates = catalog["users"]
        expected_roles = set(_role_codes_from_strings(row.roles))
        if expected_roles:
            filtered = [item for item in candidates if expected_roles.intersection({role.code for role in item.roles})]
            if filtered:
                candidates = filtered
        matched = None
        confidence = 0.0
        ambiguous = False
        if row.username:
            matches = [item for item in candidates if item.username == row.username]
            if len(matches) == 1:
                matched = matches[0]
                confidence = 0.99
            elif len(matches) > 1:
                ambiguous = True
        if not matched and row.email:
            matches = [item for item in candidates if item.email and item.email == row.email]
            if len(matches) == 1:
                matched = matches[0]
                confidence = 0.96
            elif len(matches) > 1:
                ambiguous = True
        if not matched and row.full_name:
            exact = [item for item in candidates if item.full_name.lower() == row.full_name.lower()]
            if len(exact) == 1:
                matched = exact[0]
                confidence = 0.9
            elif len(exact) > 1:
                ambiguous = True
            else:
                fuzzy = sorted(
                    ((item, _sequence_score(item.full_name, row.full_name)) for item in candidates),
                    key=lambda item: item[1],
                    reverse=True,
                )
                if fuzzy and fuzzy[0][1] >= 0.9:
                    if len(fuzzy) > 1 and abs(fuzzy[0][1] - fuzzy[1][1]) < 0.03:
                        ambiguous = True
                    else:
                        matched = fuzzy[0][0]
                        confidence = fuzzy[0][1]
        if matched:
            _apply_auto_mapping(row, matched_id=matched.id, confidence=confidence)
        elif not ambiguous and row.username and row.full_name and row.action == "unresolved":
            row.action = "create_new"
            row.confidence = 0.4

    return payload


def _week_parity_for_date(
    target_date: dt_date,
    *,
    term_start: dt_date,
    first_week_parity: Literal["odd", "even"],
) -> WeekParity:
    week_index = (target_date - term_start).days // 7
    if first_week_parity == "odd":
        return "odd" if week_index % 2 == 0 else "even"
    return "even" if week_index % 2 == 0 else "odd"


def _weekday_index(day_of_week: str) -> int:
    mapping = {
        "monday": 0,
        "tuesday": 1,
        "wednesday": 2,
        "thursday": 3,
        "friday": 4,
        "saturday": 5,
        "sunday": 6,
    }
    if day_of_week not in mapping:
        raise ValueError(f"Unsupported day_of_week '{day_of_week}'")
    return mapping[day_of_week]


def _combine_local_datetime(target_date: dt_date, time_value: str) -> datetime:
    zone = ZoneInfo(settings.app_timezone)
    return datetime.combine(target_date, _parse_time_value(time_value), tzinfo=zone).astimezone(UTC)


def _expand_pattern(pattern: DraftSchedulePatternRow, wizard: AIImportWizard) -> list[DraftLessonRow]:
    if not pattern.start_time or not pattern.end_time:
        return []
    lessons: list[DraftLessonRow] = []
    if pattern.date:
        starts_at = _combine_local_datetime(pattern.date, pattern.start_time)
        ends_at = _combine_local_datetime(pattern.date, pattern.end_time)
        lessons.append(
            DraftLessonRow(
                pattern_draft_id=pattern.draft_id,
                group_code=pattern.group_code,
                discipline_code=pattern.discipline_code,
                discipline_name=pattern.discipline_name,
                teacher_username=pattern.teacher_username,
                teacher_name=pattern.teacher_name,
                starts_at=starts_at,
                ends_at=ends_at,
                room=pattern.room,
                source_ref=pattern.source_ref,
            )
        )
        return lessons

    if not wizard.term_start or not wizard.term_end or not wizard.first_week_parity or not pattern.day_of_week:
        return []

    weekday = _weekday_index(pattern.day_of_week)
    current = wizard.term_start
    while current <= wizard.term_end:
        if current.weekday() == weekday:
            current_parity = _week_parity_for_date(
                current,
                term_start=wizard.term_start,
                first_week_parity=wizard.first_week_parity,
            )
            if pattern.week_parity == "all" or pattern.week_parity == current_parity:
                starts_at = _combine_local_datetime(current, pattern.start_time)
                ends_at = _combine_local_datetime(current, pattern.end_time)
                lessons.append(
                    DraftLessonRow(
                        pattern_draft_id=pattern.draft_id,
                        group_code=pattern.group_code,
                        discipline_code=pattern.discipline_code,
                        discipline_name=pattern.discipline_name,
                        teacher_username=pattern.teacher_username,
                        teacher_name=pattern.teacher_name,
                        starts_at=starts_at,
                        ends_at=ends_at,
                        room=pattern.room,
                        source_ref=pattern.source_ref,
                    )
                )
        current += timedelta(days=1)
    return lessons


def _build_lessons(payload: DraftPayload, wizard: AIImportWizard) -> list[DraftLessonRow]:
    lessons: list[DraftLessonRow] = []
    for pattern in payload.schedule_patterns:
        lessons.extend(_expand_pattern(pattern, wizard))
    return lessons


def _build_issues(payload: DraftPayload, wizard: AIImportWizard) -> list[DraftIssue]:
    issues: list[DraftIssue] = []
    if payload.schedule_patterns and (not wizard.term_start or not wizard.term_end or not wizard.first_week_parity):
        issues.append(
            DraftIssue(
                severity="error",
                code="wizard.calendar_required",
                message="Для расписания нужно указать период семестра и базовую чет/нечет неделю.",
                field_path="wizard",
                requires_action=True,
            )
        )

    for index, row in enumerate(payload.entities.faculties):
        if not row.code or not row.name:
            issues.append(
                DraftIssue(
                    severity="error",
                    code="faculty.required",
                    message="Для факультета нужны code и name.",
                    source_ref=row.source_ref,
                    field_path=f"entities.faculties[{index}]",
                    requires_action=True,
                )
            )
        elif row.action == "unresolved":
            issues.append(
                DraftIssue(
                    severity="warning",
                    code="faculty.mapping_unresolved",
                    message="Факультет не сопоставлен автоматически.",
                    source_ref=row.source_ref,
                    field_path=f"entities.faculties[{index}]",
                    requires_action=True,
                )
            )

    for index, row in enumerate(payload.entities.groups):
        if not row.code or not row.name:
            issues.append(
                DraftIssue(
                    severity="error",
                    code="group.required",
                    message="Для группы нужны code и name.",
                    source_ref=row.source_ref,
                    field_path=f"entities.groups[{index}]",
                    requires_action=True,
                )
            )
        elif row.action == "unresolved":
            issues.append(
                DraftIssue(
                    severity="warning",
                    code="group.mapping_unresolved",
                    message="Группа не сопоставлена автоматически.",
                    source_ref=row.source_ref,
                    field_path=f"entities.groups[{index}]",
                    requires_action=True,
                )
            )

    for index, row in enumerate(payload.entities.disciplines):
        if not row.code or not row.name:
            issues.append(
                DraftIssue(
                    severity="error",
                    code="discipline.required",
                    message="Для дисциплины нужны code и name.",
                    source_ref=row.source_ref,
                    field_path=f"entities.disciplines[{index}]",
                    requires_action=True,
                )
            )
        elif row.action == "unresolved":
            issues.append(
                DraftIssue(
                    severity="warning",
                    code="discipline.mapping_unresolved",
                    message="Дисциплина не сопоставлена автоматически.",
                    source_ref=row.source_ref,
                    field_path=f"entities.disciplines[{index}]",
                    requires_action=True,
                )
            )

    for index, row in enumerate(payload.entities.users):
        if not row.full_name:
            issues.append(
                DraftIssue(
                    severity="error",
                    code="user.full_name_required",
                    message="Для пользователя нужно ФИО.",
                    source_ref=row.source_ref,
                    field_path=f"entities.users[{index}].full_name",
                    requires_action=True,
                )
            )
        if row.action == "create_new" and not row.username:
            issues.append(
                DraftIssue(
                    severity="error",
                    code="user.username_required",
                    message="Для создания нового пользователя нужен username.",
                    source_ref=row.source_ref,
                    field_path=f"entities.users[{index}].username",
                    requires_action=True,
                )
            )
        if row.action == "unresolved":
            issues.append(
                DraftIssue(
                    severity="warning",
                    code="user.mapping_unresolved",
                    message="Пользователь не сопоставлен автоматически.",
                    source_ref=row.source_ref,
                    field_path=f"entities.users[{index}]",
                    requires_action=True,
                )
            )

    for index, row in enumerate(payload.schedule_patterns):
        if not row.group_code:
            issues.append(
                DraftIssue(
                    severity="error",
                    code="schedule.group_required",
                    message="Для шаблона расписания нужен group_code.",
                    source_ref=row.source_ref,
                    field_path=f"schedule_patterns[{index}].group_code",
                    requires_action=True,
                )
            )
        if not (row.discipline_code or row.discipline_name):
            issues.append(
                DraftIssue(
                    severity="error",
                    code="schedule.discipline_required",
                    message="Для шаблона расписания нужна дисциплина.",
                    source_ref=row.source_ref,
                    field_path=f"schedule_patterns[{index}]",
                    requires_action=True,
                )
            )
        if not (row.teacher_username or row.teacher_name):
            issues.append(
                DraftIssue(
                    severity="error",
                    code="schedule.teacher_required",
                    message="Для шаблона расписания нужен преподаватель.",
                    source_ref=row.source_ref,
                    field_path=f"schedule_patterns[{index}]",
                    requires_action=True,
                )
            )
        if not row.start_time or not row.end_time:
            issues.append(
                DraftIssue(
                    severity="error",
                    code="schedule.time_required",
                    message="Для шаблона расписания нужно время начала и конца.",
                    source_ref=row.source_ref,
                    field_path=f"schedule_patterns[{index}]",
                    requires_action=True,
                )
            )
        if not row.date and not row.day_of_week:
            issues.append(
                DraftIssue(
                    severity="error",
                    code="schedule.date_or_day_required",
                    message="Нужна дата или день недели.",
                    source_ref=row.source_ref,
                    field_path=f"schedule_patterns[{index}]",
                    requires_action=True,
                )
            )

    if payload.schedule_patterns and not payload.lessons:
        issues.append(
            DraftIssue(
                severity="warning",
                code="schedule.no_expanded_lessons",
                message="Шаблоны расписания не удалось развернуть в конкретные занятия.",
                field_path="lessons",
                requires_action=True,
            )
        )
    return issues


def _build_summary(payload: DraftPayload, issues: list[DraftIssue], extracted: ExtractedDocument | None = None) -> dict[str, Any]:
    confidences = [
        row.confidence
        for rows in (
            payload.entities.faculties,
            payload.entities.streams,
            payload.entities.groups,
            payload.entities.disciplines,
            payload.entities.users,
        )
        for row in rows
        if row.confidence is not None
    ]
    confidence = round(sum(confidences) / len(confidences), 3) if confidences else 0.0
    return {
        "detected_doc_kind": payload.detected_doc_kind,
        "confidence": confidence,
        "counts": {
            "faculties": len(payload.entities.faculties),
            "streams": len(payload.entities.streams),
            "groups": len(payload.entities.groups),
            "disciplines": len(payload.entities.disciplines),
            "users": len(payload.entities.users),
            "memberships": len(payload.entities.memberships),
            "assignments": len(payload.entities.assignments),
            "schedule_patterns": len(payload.schedule_patterns),
            "lessons": len(payload.lessons),
            "issues": len(issues),
        },
        "source_metadata": extracted.metadata if extracted else None,
        "excerpt": _ai_import_excerpt(extracted.blocks) if extracted else None,
    }


async def prepare_draft_payload(
    session: AsyncSession,
    *,
    wizard: AIImportWizard,
    payload: DraftPayload,
    extracted: ExtractedDocument | None = None,
) -> tuple[DraftPayload, list[DraftIssue], dict[str, Any]]:
    payload = DraftPayload.model_validate(payload.model_dump())
    payload = _dedupe_payload(payload)
    _derive_memberships_and_assignments(payload)
    await _match_payload(session, payload)
    payload.lessons = _build_lessons(payload, wizard)
    issues = _build_issues(payload, wizard)
    summary = _build_summary(payload, issues, extracted)
    return payload, issues, summary


async def build_ai_import_draft_payload(
    session: AsyncSession,
    *,
    mode: AIImportMode,
    wizard: AIImportWizard,
    file_path: str,
) -> tuple[DraftPayload, list[DraftIssue], dict[str, Any]]:
    extracted = await asyncio.to_thread(extract_document, file_path)
    if not extracted.blocks:
        raise RuntimeError("Could not extract any text or table data from the document")
    parsed = await _normalize_with_llm(mode, extracted)
    payload = _parsed_to_payload(parsed)
    return await prepare_draft_payload(session, wizard=wizard, payload=payload, extracted=extracted)


def _serialize_issue(issue: DraftIssue) -> dict[str, Any]:
    return issue.model_dump(mode="json")


def serialize_ai_import_draft(row: AIImportDraft) -> dict[str, Any]:
    return {
        "id": row.id,
        "status": row.status.value,
        "mode": row.mode.value,
        "file_name": row.file_name,
        "created_at": row.created_at,
        "updated_at": row.updated_at,
        "completed_at": row.completed_at,
        "wizard": row.wizard or {},
        "summary": row.summary,
        "payload": row.payload,
        "issues": row.issues or [],
        "apply_result": row.apply_result,
        "error_report": row.error_report,
    }


async def mark_ai_import_draft_processing(session: AsyncSession, draft: AIImportDraft) -> None:
    draft.status = AIImportDraftStatus.PROCESSING
    draft.error_report = None
    await session.commit()


async def mark_ai_import_draft_failed(session: AsyncSession, draft: AIImportDraft, detail: str) -> None:
    draft.status = AIImportDraftStatus.FAILED
    draft.error_report = {"error": detail}
    draft.completed_at = utc_now()
    await session.commit()


async def process_ai_import_draft_record(session: AsyncSession, draft: AIImportDraft) -> None:
    wizard = AIImportWizard.model_validate(draft.wizard or {})
    payload, issues, summary = await build_ai_import_draft_payload(
        session,
        mode=draft.mode,
        wizard=wizard,
        file_path=draft.file_path,
    )
    draft.status = AIImportDraftStatus.DRAFT
    draft.payload = payload.model_dump(mode="json")
    draft.issues = [_serialize_issue(item) for item in issues]
    draft.summary = summary
    draft.completed_at = utc_now()
    draft.error_report = None
    await session.commit()


async def update_ai_import_draft_payload(
    session: AsyncSession,
    *,
    draft: AIImportDraft,
    wizard: AIImportWizard,
    payload_data: dict[str, Any],
    actor_user_id: UUID,
) -> AIImportDraft:
    if draft.status not in {AIImportDraftStatus.DRAFT, AIImportDraftStatus.FAILED}:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Draft cannot be edited in current status")
    payload = DraftPayload.model_validate(payload_data)
    payload, issues, summary = await prepare_draft_payload(session, wizard=wizard, payload=payload)
    draft.wizard = wizard.model_dump(mode="json")
    draft.payload = payload.model_dump(mode="json")
    draft.issues = [_serialize_issue(item) for item in issues]
    draft.summary = summary
    draft.status = AIImportDraftStatus.DRAFT
    draft.error_report = None
    await log_audit(
        session,
        actor_user_id=actor_user_id,
        action="ai_import_updated",
        entity_type="ai_import_draft",
        entity_id=str(draft.id),
        details={"mode": draft.mode.value},
    )
    await session.commit()
    await session.refresh(draft)
    return draft


def _resolve_group_row(payload: DraftPayload, group_code: str | None) -> DraftGroupRow | None:
    if not group_code:
        return None
    return next((row for row in payload.entities.groups if row.code == group_code), None)


def _resolve_discipline_row(
    payload: DraftPayload,
    *,
    discipline_code: str | None,
    discipline_name: str | None = None,
) -> DraftDisciplineRow | None:
    if discipline_code:
        match = next((row for row in payload.entities.disciplines if row.code == discipline_code), None)
        if match:
            return match
    if discipline_name:
        return next((row for row in payload.entities.disciplines if row.name == discipline_name), None)
    return None


def _resolve_user_row(
    payload: DraftPayload,
    *,
    username: str | None,
    full_name: str | None = None,
) -> DraftUserRow | None:
    if username:
        match = next((row for row in payload.entities.users if row.username == username), None)
        if match:
            return match
    if full_name:
        return next((row for row in payload.entities.users if row.full_name == full_name), None)
    return None


def _issues_require_action(issues: list[DraftIssue]) -> bool:
    return any(issue.requires_action or issue.severity == "error" for issue in issues)


async def apply_ai_import_draft(
    session: AsyncSession,
    *,
    draft: AIImportDraft,
    actor_user_id: UUID,
) -> dict[str, Any]:
    if draft.status != AIImportDraftStatus.DRAFT:
        raise HTTPException(status_code=status.HTTP_409_CONFLICT, detail="Only draft imports can be applied")
    wizard = AIImportWizard.model_validate(draft.wizard or {})
    payload = DraftPayload.model_validate(draft.payload or {})
    payload, issues, summary = await prepare_draft_payload(session, wizard=wizard, payload=payload)
    if _issues_require_action(issues):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Resolve draft issues before apply")

    faculty_refs: dict[str, Faculty] = {}
    stream_refs: dict[str, Stream] = {}
    group_refs: dict[str, Group] = {}
    discipline_refs: dict[str, Discipline] = {}
    user_refs: dict[str, User] = {}

    try:
        for row in payload.entities.faculties:
            faculty_refs[row.draft_id] = await resolve_faculty(
                session,
                code=row.code,
                name=row.name,
                action="match_existing" if row.action == "match_existing" else "create_new",
                existing_id=row.existing_id,
            )

        for row in payload.entities.streams:
            faculty_id = None
            faculty_row = next(
                (item for item in payload.entities.faculties if item.code == row.faculty_code),
                None,
            )
            if faculty_row and faculty_row.draft_id in faculty_refs:
                faculty_id = faculty_refs[faculty_row.draft_id].id
            stream_refs[row.draft_id] = await resolve_stream(
                session,
                name=row.name,
                faculty_id=faculty_id,
                action="match_existing" if row.action == "match_existing" else "create_new",
                existing_id=row.existing_id,
            )

        for row in payload.entities.groups:
            faculty_id = None
            stream_id = None
            parent_group_id = None
            faculty_row = next(
                (item for item in payload.entities.faculties if item.code == row.faculty_code),
                None,
            )
            if faculty_row and faculty_row.draft_id in faculty_refs:
                faculty_id = faculty_refs[faculty_row.draft_id].id
            stream_row = next(
                (item for item in payload.entities.streams if item.name == row.stream_name and item.faculty_code == row.faculty_code),
                None,
            )
            if stream_row and stream_row.draft_id in stream_refs:
                stream_id = stream_refs[stream_row.draft_id].id
            parent_group = next((item for item in payload.entities.groups if item.code == row.parent_group_code), None)
            if parent_group and parent_group.draft_id in group_refs:
                parent_group_id = group_refs[parent_group.draft_id].id
            group_refs[row.draft_id] = await resolve_group(
                session,
                code=row.code,
                name=row.name,
                faculty_id=faculty_id,
                stream_id=stream_id,
                parent_group_id=parent_group_id,
                is_subgroup=row.is_subgroup,
                action="match_existing" if row.action == "match_existing" else "create_new",
                existing_id=row.existing_id,
            )

        for row in payload.entities.disciplines:
            discipline_refs[row.draft_id] = await resolve_discipline(
                session,
                code=row.code,
                name=row.name,
                action="match_existing" if row.action == "match_existing" else "create_new",
                existing_id=row.existing_id,
            )

        for row in payload.entities.users:
            user_refs[row.draft_id] = await resolve_user(
                session,
                username=row.username,
                full_name=row.full_name,
                email=row.email,
                role_codes=_role_codes_from_strings(row.roles),
                action="match_existing" if row.action == "match_existing" else "create_new",
                existing_id=row.existing_id,
                role_update_strategy="merge",
            )

        membership_count = 0
        assignment_count = 0
        lesson_count = 0

        for row in payload.entities.memberships:
            user_row = _resolve_user_row(
                payload,
                username=row.student_username,
                full_name=row.student_full_name,
            )
            group_row = _resolve_group_row(payload, row.group_code)
            if not user_row or not group_row:
                continue
            await ensure_student_membership(
                session,
                student_id=user_refs[user_row.draft_id].id,
                group_id=group_refs[group_row.draft_id].id,
                start_date=row.start_date or utc_now().date(),
            )
            membership_count += 1

        for row in payload.entities.assignments:
            user_row = _resolve_user_row(
                payload,
                username=row.teacher_username,
                full_name=row.teacher_full_name,
            )
            group_row = _resolve_group_row(payload, row.group_code)
            discipline_row = _resolve_discipline_row(
                payload,
                discipline_code=row.discipline_code,
                discipline_name=row.discipline_name,
            )
            if not user_row or not group_row or not discipline_row:
                continue
            await ensure_teacher_assignment(
                session,
                teacher_id=user_refs[user_row.draft_id].id,
                discipline_id=discipline_refs[discipline_row.draft_id].id,
                group_id=group_refs[group_row.draft_id].id,
            )
            assignment_count += 1

        for row in payload.lessons:
            group_row = _resolve_group_row(payload, row.group_code)
            discipline_row = _resolve_discipline_row(
                payload,
                discipline_code=row.discipline_code,
                discipline_name=row.discipline_name,
            )
            user_row = _resolve_user_row(payload, username=row.teacher_username, full_name=row.teacher_name)
            if not group_row or not discipline_row or not user_row:
                continue
            await upsert_lesson(
                session,
                group=group_refs[group_row.draft_id],
                discipline=discipline_refs[discipline_row.draft_id],
                teacher=user_refs[user_row.draft_id],
                starts_at=row.starts_at,
                ends_at=row.ends_at,
                room=row.room,
                status=LessonStatus.PLANNED,
            )
            lesson_count += 1

        draft.status = AIImportDraftStatus.APPLIED
        draft.payload = payload.model_dump(mode="json")
        draft.issues = [_serialize_issue(item) for item in issues]
        draft.summary = summary
        draft.apply_result = {
            "faculties": len(payload.entities.faculties),
            "streams": len(payload.entities.streams),
            "groups": len(payload.entities.groups),
            "disciplines": len(payload.entities.disciplines),
            "users": len(payload.entities.users),
            "memberships": membership_count,
            "assignments": assignment_count,
            "lessons": lesson_count,
        }
        draft.completed_at = utc_now()
        draft.error_report = None
        await log_audit(
            session,
            actor_user_id=actor_user_id,
            action="ai_import_applied",
            entity_type="ai_import_draft",
            entity_id=str(draft.id),
            details={"mode": draft.mode.value, "lessons": lesson_count},
        )
        await session.commit()
        return draft.apply_result or {}
    except HTTPException:
        await session.rollback()
        raise
    except ValueError as exc:
        await session.rollback()
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(exc)) from exc
    except Exception as exc:
        await session.rollback()
        logger.exception("ai_import_apply_failed", draft_id=str(draft.id), reason=str(exc))
        raise HTTPException(status_code=status.HTTP_500_INTERNAL_SERVER_ERROR, detail="AI import apply failed") from exc


async def reject_ai_import_draft(
    session: AsyncSession,
    *,
    draft: AIImportDraft,
    actor_user_id: UUID,
) -> None:
    draft.status = AIImportDraftStatus.REJECTED
    draft.completed_at = utc_now()
    await log_audit(
        session,
        actor_user_id=actor_user_id,
        action="ai_import_rejected",
        entity_type="ai_import_draft",
        entity_id=str(draft.id),
        details={"mode": draft.mode.value},
    )
    await session.commit()
