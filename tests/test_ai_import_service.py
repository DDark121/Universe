from __future__ import annotations

from pathlib import Path
from uuid import UUID

import pytest
from docx import Document as DocxDocument
from httpx import ASGITransport, AsyncClient
from sqlalchemy import select

from app.api.deps import get_current_user
from app.core.db import get_db_session
from app.db.enums import AIImportDraftStatus, RoleCode
from app.db.models import (
    AIImportDraft,
    Discipline,
    Group,
    Lesson,
    Role,
    StudentGroupMembership,
    TeacherAssignment,
    User,
)
from app.main import app
from app.services.ai_imports import (
    AIImportWizard,
    DraftEntities,
    DraftPayload,
    DraftSchedulePatternRow,
    ParsedDiscipline,
    ParsedDocument,
    ParsedGroup,
    ParsedUser,
    _extract_docx_document,
    _extract_pdf_document,
    prepare_draft_payload,
    process_ai_import_draft_record,
)


@pytest.fixture()
async def api_client(session):
    async def override_db():
        yield session

    app.dependency_overrides[get_db_session] = override_db
    async with AsyncClient(transport=ASGITransport(app=app), base_url="http://test") as client:
        yield client
    app.dependency_overrides.clear()


def override_user(user: User) -> None:
    async def _override_user():
        return user

    app.dependency_overrides[get_current_user] = _override_user


def build_simple_pdf(path: Path, text: str) -> None:
    stream = f"BT\n/F1 18 Tf\n36 120 Td\n({text}) Tj\nET"
    content = stream.encode("latin-1")
    objects = [
        b"1 0 obj<< /Type /Catalog /Pages 2 0 R >>endobj\n",
        b"2 0 obj<< /Type /Pages /Kids [3 0 R] /Count 1 >>endobj\n",
        b"3 0 obj<< /Type /Page /Parent 2 0 R /MediaBox [0 0 300 180] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>endobj\n",
        f"4 0 obj<< /Length {len(content)} >>stream\n".encode("latin-1") + content + b"\nendstream\nendobj\n",
        b"5 0 obj<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>endobj\n",
    ]
    header = b"%PDF-1.4\n"
    body = bytearray(header)
    offsets = [0]
    for obj in objects:
        offsets.append(len(body))
        body.extend(obj)
    xref_offset = len(body)
    body.extend(f"xref\n0 {len(objects) + 1}\n".encode("latin-1"))
    body.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        body.extend(f"{offset:010d} 00000 n \n".encode("latin-1"))
    body.extend(
        f"trailer<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode("latin-1")
    )
    path.write_bytes(bytes(body))


@pytest.mark.asyncio
async def test_extract_docx_and_pdf_documents(tmp_path: Path):
    docx_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_paragraph("Группа SE-101")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Преподаватель"
    table.cell(0, 1).text = "Дисциплина"
    table.cell(1, 0).text = "teacher_math"
    table.cell(1, 1).text = "Math"
    document.save(docx_path)

    pdf_path = tmp_path / "sample.pdf"
    build_simple_pdf(pdf_path, "SE-101 Math")

    docx_result = _extract_docx_document(str(docx_path))
    pdf_result = _extract_pdf_document(str(pdf_path))

    assert any("Группа SE-101" in block.text for block in docx_result.blocks)
    assert any("teacher_math" in block.text for block in docx_result.blocks)
    assert any("SE-101 Math" in block.text for block in pdf_result.blocks)


@pytest.mark.asyncio
async def test_prepare_draft_payload_expands_semester_schedule(session):
    payload = DraftPayload(
        detected_doc_kind="schedule",
        entities=DraftEntities(
            groups=[],
            disciplines=[],
            users=[],
        ),
        schedule_patterns=[
            DraftSchedulePatternRow(
                group_code="SE-101",
                discipline_code="MATH",
                teacher_username="teacher_math",
                day_of_week="monday",
                start_time="08:30",
                end_time="10:00",
                week_parity="odd",
            )
        ],
    )

    prepared, issues, summary = await prepare_draft_payload(
        session,
        wizard=AIImportWizard(term_start="2026-09-07", term_end="2026-09-21", first_week_parity="odd"),
        payload=payload,
    )

    assert summary["counts"]["lessons"] == 2
    assert not any(issue.code == "wizard.calendar_required" for issue in issues)
    assert [lesson.starts_at.date().isoformat() for lesson in prepared.lessons] == ["2026-09-07", "2026-09-21"]


@pytest.mark.asyncio
async def test_admin_ai_import_flow_creates_entities_and_lessons(session, api_client, monkeypatch, tmp_path: Path):
    admin_role = Role(code=RoleCode.ADMIN, name="Admin")
    teacher_role = Role(code=RoleCode.TEACHER, name="Teacher")
    student_role = Role(code=RoleCode.STUDENT, name="Student")
    admin = User(username="admin_ai", full_name="Admin AI", password_hash="x", must_change_password=False)
    admin.roles.append(admin_role)
    session.add_all([admin_role, teacher_role, student_role, admin])
    await session.commit()
    override_user(admin)

    monkeypatch.setattr("app.api.v1.admin.process_ai_import_draft.delay", lambda _draft_id: None)
    monkeypatch.setattr("app.services.ai_imports.settings.ai_imports_dir", str(tmp_path / "ai-imports"))

    async def fake_normalize(_mode, _extracted):
        return ParsedDocument(
            detected_doc_kind="mixed",
            groups=[ParsedGroup(code="SE-101", name="SE-101")],
            disciplines=[ParsedDiscipline(code="MATH", name="Math")],
            users=[
                ParsedUser(username="teacher_math", full_name="Teacher Math", roles=["teacher"]),
                ParsedUser(username="student_1", full_name="Student One", roles=["student"], group_code="SE-101"),
            ],
            schedule_patterns=[
                {
                    "group_code": "SE-101",
                    "discipline_code": "MATH",
                    "teacher_username": "teacher_math",
                    "day_of_week": "monday",
                    "start_time": "08:30",
                    "end_time": "10:00",
                    "week_parity": "all",
                }
            ],
        )

    monkeypatch.setattr("app.services.ai_imports._normalize_with_llm", fake_normalize)

    response = await api_client.post(
        "/api/v1/admin/ai-imports",
        data={
            "mode": "mixed",
            "term_start": "2026-09-07",
            "term_end": "2026-09-14",
            "first_week_parity": "odd",
        },
        files={"file": ("schedule.csv", b"group,teacher\nSE-101,teacher_math\n", "text/csv")},
    )
    assert response.status_code == 200
    draft_id = response.json()["id"]

    draft = (await session.execute(select(AIImportDraft).where(AIImportDraft.id == UUID(draft_id)))).scalar_one()
    await process_ai_import_draft_record(session, draft)
    await session.refresh(draft)
    assert draft.status == AIImportDraftStatus.DRAFT
    assert draft.payload is not None
    assert Path(draft.file_path).exists()
    assert Path(draft.file_path).is_relative_to(tmp_path / "ai-imports")

    apply_response = await api_client.post(f"/api/v1/admin/ai-imports/{draft_id}/apply")
    assert apply_response.status_code == 200

    created_group = (await session.execute(select(Group).where(Group.code == "SE-101"))).scalar_one()
    created_discipline = (await session.execute(select(Discipline).where(Discipline.code == "MATH"))).scalar_one()
    created_teacher = (await session.execute(select(User).where(User.username == "teacher_math"))).scalar_one()
    created_student = (await session.execute(select(User).where(User.username == "student_1"))).scalar_one()
    assignment = (
        await session.execute(
            select(TeacherAssignment).where(
                TeacherAssignment.teacher_id == created_teacher.id,
                TeacherAssignment.discipline_id == created_discipline.id,
                TeacherAssignment.group_id == created_group.id,
            )
        )
    ).scalar_one()
    membership = (
        await session.execute(
            select(StudentGroupMembership).where(
                StudentGroupMembership.student_id == created_student.id,
                StudentGroupMembership.group_id == created_group.id,
            )
        )
    ).scalar_one()
    lessons = (
        await session.execute(
            select(Lesson).where(
                Lesson.group_id == created_group.id,
                Lesson.discipline_id == created_discipline.id,
                Lesson.teacher_id == created_teacher.id,
            )
        )
    ).scalars().all()

    assert assignment.is_active is True
    assert membership.is_primary is True
    assert len(lessons) == 2
